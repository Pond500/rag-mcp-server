"""
Multi-Knowledge Base RAG Pipeline
Hybrid Approach: Simple but Powerful

Features:
- Dynamic collection creation (auto-create if not exists)
- Upload to new or existing collections
- Query single or multiple collections
- Collection management (list, info, delete)
"""

import qdrant_client
from qdrant_client.models import Distance, VectorParams, PointStruct
from typing import List, Dict, Any, Optional, Tuple
import uuid
from datetime import datetime
import json
from app.prompts import METADATA_EXTRACTOR_TEMPLATE
from app.logger import logger

# LangChain imports
from langchain_core.documents import Document
from langchain_openai import ChatOpenAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_qdrant import Qdrant
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory

# Local imports
import app.config as config
from app.ocr_service import get_text_from_pdf
import io


class MultiKnowledgeBaseRAG:
    """Multi-Knowledge Base RAG System with Hybrid Approach"""
    
    # Router Index constant
    ROUTER_COLLECTION_NAME = "master_router_index"
    ROUTER_SIMILARITY_THRESHOLD = 0.4  # Minimum score for routing confidence
    
    def __init__(self):
        """Initialize Multi-KB RAG system"""
        logger.info("🚀 Initializing Multi-Knowledge Base RAG System...")
        
        # Qdrant client
        self.qdrant_client = qdrant_client.QdrantClient(
            host=config.QDRANT_HOST,
            port=config.QDRANT_PORT
        )
        logger.debug(f"Qdrant client: {config.QDRANT_HOST}:{config.QDRANT_PORT}")
        
        # LLM
        self.llm = ChatOpenAI(
            model=config.LLM_MODEL_NAME,
            base_url=config.LLM_API_BASE,
            api_key=config.LLM_API_KEY,
            max_tokens=1500,
            temperature=0.5,
            request_timeout=120
        )
        
        # Embedding model
        self.embed_model = HuggingFaceEmbeddings(
            model_name=config.EMBED_MODEL_NAME,
            encode_kwargs={'normalize_embeddings': True}
        )
        
        # Text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP
        )
        
        # Chat histories: {collection_name: {session_id: ConversationBufferMemory}}
        self.chat_histories: Dict[str, Dict[str, ConversationBufferMemory]] = {}
        
        # Ensure router index exists
        self._ensure_router_index()
        
        logger.info("✅ Multi-KB RAG System initialized")
    
    def _normalize_collection_name(self, name: str) -> str:
        """Normalize collection name (lowercase, replace spaces with underscore)"""
        return name.lower().replace(" ", "_").replace("-", "_")
    
    def _get_collection_name(self, kb_name: str) -> str:
        """Get full collection name with prefix"""
        normalized = self._normalize_collection_name(kb_name)
        return f"kb_{normalized}"
    
    def _ensure_router_index(self) -> None:
        """Ensure master router index exists (for semantic routing)"""
        if not self.collection_exists(self.ROUTER_COLLECTION_NAME):
            try:
                # Create router index with same dimensions as embeddings (1024 for bge-m3)
                self.qdrant_client.create_collection(
                    collection_name=self.ROUTER_COLLECTION_NAME,
                    vectors_config=VectorParams(size=1024, distance=Distance.COSINE)
                )
                logger.info(f"✅ Created master router index: {self.ROUTER_COLLECTION_NAME}")
            except Exception as e:
                logger.error(f"⚠️ Failed to create router index: {e}", exc_info=True)
    
    def _update_router_index(self, kb_name: str, description: str) -> None:
        """
        Update master router index with KB description
        This allows semantic routing to find the right KB based on query similarity
        
        Args:
            kb_name: Knowledge base name
            description: KB description (AI-generated or manual)
        """
        if not description or description == "Auto-created collection":
            logger.debug(f"Skipping router update for {kb_name} (no meaningful description)")
            return
        
        try:
            logger.debug(f"Updating router index for {kb_name}...")
            # Generate embedding for the description
            description_vector = self.embed_model.embed_query(description)
            
            # Create point for router index
            collection_name = self._get_collection_name(kb_name)
            point = PointStruct(
                id=collection_name,  # Use collection name as ID for easy updates
                vector=description_vector,
                payload={
                    "kb_name": kb_name,
                    "collection_name": collection_name,
                    "description": description,
                    "updated_at": datetime.now().isoformat()
                }
            )
            
            # Upsert to router index
            self.qdrant_client.upsert(
                collection_name=self.ROUTER_COLLECTION_NAME,
                points=[point]
            )
            
            logger.info(f"✅ Updated router index for: {kb_name}")
            logger.debug(f"Description: {description[:100]}...")
            
        except Exception as e:
            logger.error(f"Failed to update router index for {kb_name}: {e}", exc_info=True)
    
    def collection_exists(self, collection_name: str) -> bool:
        """Check if collection exists"""
        try:
            self.qdrant_client.get_collection(collection_name)
            return True
        except Exception:
            return False
    
    def create_collection(self, kb_name: str, description: str = "") -> Dict[str, Any]:
        """
        Create new knowledge base (collection)
        
        Args:
            kb_name: Knowledge base name (e.g., "medical", "legal")
            description: Optional description
            
        Returns:
            Dict with collection info
        """
        collection_name = self._get_collection_name(kb_name)
        
        if self.collection_exists(collection_name):
            return {
                "success": False,
                "message": f"Collection '{kb_name}' already exists",
                "collection_name": collection_name
            }
        
        try:
            # Create collection (bge-m3 dimension = 1024)
            self.qdrant_client.create_collection(
                collection_name=collection_name,
                vectors_config=VectorParams(size=1024, distance=Distance.COSINE)
            )
            
            # Store metadata as payload in a special point
            metadata_point = PointStruct(
                id=str(uuid.uuid4()),
                vector=[0.0] * 1024,  # Dummy vector
                payload={
                    "_type": "collection_metadata",
                    "kb_name": kb_name,
                    "description": description,
                    "created_at": datetime.now().isoformat(),
                    "document_count": 0
                }
            )
            self.qdrant_client.upsert(
                collection_name=collection_name,
                points=[metadata_point]
            )
            
            # Update router index for semantic routing
            self._update_router_index(kb_name, description)
            
            logger.info(f"✅ Created collection: {collection_name}")
            return {
                "success": True,
                "kb_name": kb_name,
                "collection_name": collection_name,
                "description": description,
                "created_at": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"Failed to create collection: {e}", exc_info=True)
            return {
                "success": False,
                "message": str(e)
            }
    
    def list_collections(self) -> Dict[str, Any]:
        """
        List all knowledge base collections with descriptions
        
        Returns:
            Dict with collections list and count
        """
        try:
            collections = self.qdrant_client.get_collections().collections
            kb_list = []
            
            for col in collections:
                if col.name.startswith("kb_"):
                    kb_name = col.name[3:]  # Remove "kb_" prefix
                    
                    # Get collection info
                    info = self.qdrant_client.get_collection(col.name)
                    
                    # Try to get metadata (including AI-generated description)
                    description = "No description"
                    try:
                        results = self.qdrant_client.scroll(
                            collection_name=col.name,
                            scroll_filter={
                                "must": [
                                    {"key": "_type", "match": {"value": "collection_metadata"}}
                                ]
                            },
                            limit=1
                        )
                        if results[0]:
                            metadata_payload = results[0][0].payload
                            description = metadata_payload.get("description", "No description")
                    except Exception:
                        pass
                    
                    kb_list.append({
                        "kb_name": kb_name,
                        "collection_name": col.name,
                        "description": description,
                        "points_count": info.points_count,
                        "vectors_count": info.vectors_count
                    })
            
            return {
                "success": True,
                "count": len(kb_list),
                "collections": kb_list
            }
        except Exception as e:
            logger.error(f"Failed to list collections: {e}", exc_info=True)
            return {
                "success": False,
                "message": str(e),
                "collections": []
            }
    
    def get_collection_info(self, kb_name: str) -> Optional[Dict[str, Any]]:
        """Get detailed info about a collection"""
        collection_name = self._get_collection_name(kb_name)
        
        if not self.collection_exists(collection_name):
            return None
        
        try:
            info = self.qdrant_client.get_collection(collection_name)
            
            # Try to get metadata
            metadata = None
            try:
                results = self.qdrant_client.scroll(
                    collection_name=collection_name,
                    scroll_filter={
                        "must": [
                            {"key": "_type", "match": {"value": "collection_metadata"}}
                        ]
                    },
                    limit=1
                )
                if results[0]:
                    metadata = results[0][0].payload
            except Exception:
                pass
            
            return {
                "kb_name": kb_name,
                "collection_name": collection_name,
                "points_count": info.points_count,
                "vectors_count": info.vectors_count,
                "metadata": metadata
            }
        except Exception as e:
            logger.error(f"Failed to get collection info: {e}", exc_info=True)
            return None
    
    def delete_collection(self, kb_name: str) -> Dict[str, Any]:
        """Delete a knowledge base collection"""
        collection_name = self._get_collection_name(kb_name)
        
        if not self.collection_exists(collection_name):
            return {
                "success": False,
                "message": f"Collection '{kb_name}' not found"
            }
        
        try:
            self.qdrant_client.delete_collection(collection_name)
            
            # Clear chat histories for this collection
            if collection_name in self.chat_histories:
                del self.chat_histories[collection_name]
            
            logger.info(f"✅ Deleted collection: {collection_name}")
            return {
                "success": True,
                "message": f"Collection '{kb_name}' deleted successfully"
            }
        except Exception as e:
            logger.error(f"Failed to delete collection: {e}", exc_info=True)
            return {
                "success": False,
                "message": str(e)
            }
    
    def _extract_text_from_file(self, file_bytes: bytes, content_type: str) -> List[Dict[str, Any]]:
        """Extract text from file based on content type"""
        if content_type == "application/pdf":
            return get_text_from_pdf(file_bytes)
        
        elif content_type == "text/plain":
            try:
                text = file_bytes.decode('utf-8')
                return [{"page_number": 1, "text": text}]
            except Exception as e:
                logger.error(f"Failed to read TXT: {e}", exc_info=True)
                return []
        
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                text = '\n'.join(full_text)
                return [{"page_number": 1, "text": text}]
            except Exception as e:
                logger.error(f"Failed to read DOCX: {e}", exc_info=True)
                return []
        
        else:
            logger.warning(f"Unsupported file type: {content_type}")
            return []
    
    def _extract_metadata_from_text(self, text: str) -> Dict[str, Any]:
        """
        Extract metadata from document text using LLM
        
        Args:
            text: Document text (usually first page)
            
        Returns:
            Dict with doc_type, category, status, title
        """
        logger.info("🤖 AI Extracting metadata from document...")
        
        # Truncate text to save tokens (first 4000 characters)
        truncated_text = text[:4000]
        
        try:
            # Create prompt using template from app.prompts
            prompt_str = METADATA_EXTRACTOR_TEMPLATE.format(context_str=truncated_text)
            
            # Call LLM
            response = self.llm.invoke(prompt_str)
            raw_output = response.content
            
            # Clean and parse JSON output
            json_str = raw_output.strip()
            # Remove markdown code blocks if present
            if json_str.startswith("```"):
                json_str = json_str.split("```")[1]
                if json_str.startswith("json"):
                    json_str = json_str[4:]
            json_str = json_str.strip()
            
            extracted_data = json.loads(json_str)
            
            # Validate required fields
            if isinstance(extracted_data, dict):
                required_fields = ["doc_type", "category", "status", "title"]
                for field in required_fields:
                    if field not in extracted_data:
                        extracted_data[field] = "Unknown"
                
                logger.debug(f"AI Metadata extracted: {extracted_data}")
                return extracted_data
            else:
                raise ValueError("LLM did not return a dictionary")
            
        except Exception as e:
            logger.warning(f"Metadata extraction failed: {e}")
            # Fallback to default values
            return {
                "doc_type": "Unknown", 
                "category": "General", 
                "status": "Published", 
                "title": "Untitled Document"
            }
    
    def _generate_smart_description(self, metadata: Dict[str, Any], filename: str) -> str:
        """
        Generate smart collection description from AI-extracted metadata
        
        Args:
            metadata: Metadata dictionary from _extract_metadata_from_text
            filename: Original filename
            
        Returns:
            Rich description string for Semantic Router
        """
        doc_type = metadata.get("doc_type", "Unknown")
        title = metadata.get("title", filename)
        category = metadata.get("category", "General")
        status = metadata.get("status", "")
        
        # Format: "[Type] Title - Category: Category (Status)"
        description = f"[{doc_type}] {title} - Category: {category}"
        if status and status != "Unknown":
            description += f" ({status})"
        
        return description

    def upload_document(
        self,
        kb_name: str,
        file_bytes: bytes,
        filename: str,
        content_type: str,
        auto_create: bool = True,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Upload document to knowledge base with AI-First Ingestion Flow
        
        NEW FLOW (Semantic Router Ready):
        1. Extract text from file FIRST
        2. Use AI to extract metadata (doc_type, title, category, status)
        3. Generate smart description for collection
        4. Create collection with rich description (if auto_create=True)
        5. Store document with AI metadata
        
        Args:
            kb_name: Knowledge base name
            file_bytes: File content
            filename: File name
            content_type: MIME type
            auto_create: Auto-create collection if not exists (default: True)
            metadata: Additional metadata (optional, will override AI metadata)
            
        Returns:
            Dict with upload result including AI-extracted metadata
        """
        collection_name = self._get_collection_name(kb_name)
        
        # ========================================
        # STEP 1: Extract text from file FIRST
        # ========================================
        logger.info(f"📄 Extracting text from {filename}...")
        page_data_list = self._extract_text_from_file(file_bytes, content_type)
        
        if not page_data_list:
            return {
                "success": False,
                "message": "Failed to extract text from file. Please check file format and content."
            }
        
        # ========================================
        # STEP 2: Use AI to extract metadata from first page
        # ========================================
        ai_metadata = {}
        smart_description = f"Auto-created collection for {filename}"
        
        if page_data_list and page_data_list[0].get('text'):
            first_page_text = page_data_list[0]['text']
            ai_metadata = self._extract_metadata_from_text(first_page_text)
            
            # ========================================
            # STEP 3: Generate smart description for Semantic Router
            # ========================================
            smart_description = self._generate_smart_description(ai_metadata, filename)
            logger.info(f"📝 Generated description: {smart_description}")
        else:
            logger.warning("No text found in first page, skipping AI metadata extraction")
        
        # ========================================
        # STEP 4: Create collection with rich description (if needed)
        # ========================================
        if not self.collection_exists(collection_name):
            if auto_create:
                logger.info(f"📦 Collection '{kb_name}' not found, creating with AI-generated description...")
                result = self.create_collection(kb_name, description=smart_description)
                if not result["success"]:
                    return result
                logger.info(f"✅ Collection created: {collection_name}")
                logger.debug(f"Description: {smart_description}")
            else:
                return {
                    "success": False,
                    "message": f"Knowledge base '{kb_name}' does not exist.",
                    "suggestion": f"Set auto_create=true to automatically create the KB, or call create_collection(kb_name='{kb_name}') first.",
                    "ai_metadata": ai_metadata  # Return AI findings even on error
                }
        
        # ========================================
        # STEP 5: Create documents with AI metadata
        # ========================================
        logger.info(f"📝 Creating documents from {len(page_data_list)} pages...")
        documents: List[Document] = []
        
        for page_data in page_data_list:
            # Merge metadata: System + AI + User-provided (priority order)
            doc_metadata = {
                "kb_name": kb_name,
                "filename": filename,
                "page_number": page_data['page_number'],
                "uploaded_at": datetime.now().isoformat(),
                **ai_metadata,       # AI-extracted metadata
                **(metadata or {})   # User metadata (overrides AI if provided)
            }
            
            doc = Document(
                page_content=page_data['text'],
                metadata=doc_metadata
            )
            documents.append(doc)
        
        # Split documents into chunks
        logger.info(f"✂️ Splitting documents into chunks...")
        split_docs = self.text_splitter.split_documents(documents)
        logger.debug(f"Created {len(split_docs)} chunks")
        
        # ========================================
        # STEP 6: Store in Qdrant
        # ========================================
        try:
            vector_store = Qdrant(
                client=self.qdrant_client,
                collection_name=collection_name,
                embeddings=self.embed_model
            )
            vector_store.add_documents(split_docs)
            
            # ========================================
            # STEP 7: Update router index with smart description
            # ========================================
            self._update_router_index(kb_name, smart_description)
            
            logger.info(f"✅ Successfully uploaded {filename} to {kb_name}")
            logger.debug(f"KB Name: {kb_name}")
            logger.debug(f"Collection: {collection_name}")
            logger.debug(f"Description: {smart_description}")
            logger.debug(f"AI Metadata: {ai_metadata}")
            
            return {
                "success": True,
                "kb_name": kb_name,
                "collection_name": collection_name,
                "filename": filename,
                "chunks": len(split_docs),
                "pages": len(page_data_list),
                "ai_metadata": ai_metadata,
                "collection_description": smart_description,
                "message": f"Document uploaded successfully to '{kb_name}' with AI-generated metadata"
            }
        except Exception as e:
            logger.error(f"Failed to upload document: {e}", exc_info=True)
            import traceback
            pass  # traceback handled by logger
            return {
                "success": False,
                "message": f"Failed to store document in Qdrant: {str(e)}",
                "ai_metadata": ai_metadata  # Return AI findings even on error
            }
    
    def _get_or_create_memory(self, collection_name: str, session_id: str) -> ConversationBufferMemory:
        """Get or create conversation memory for a session"""
        if collection_name not in self.chat_histories:
            self.chat_histories[collection_name] = {}
        
        if session_id not in self.chat_histories[collection_name]:
            self.chat_histories[collection_name][session_id] = ConversationBufferMemory(
                memory_key="chat_history",
                return_messages=True,
                output_key="answer"
            )
        
        return self.chat_histories[collection_name][session_id]
    
    def chat_with_collection(
        self,
        kb_name: str,
        query: str,
        session_id: str,
        top_k: int = 5
    ) -> Dict[str, Any]:
        """
        Chat with a single knowledge base
        
        Args:
            kb_name: Knowledge base name
            query: User query
            session_id: Session ID for conversation history
            top_k: Number of documents to retrieve
            
        Returns:
            Dict with answer and sources
        """
        collection_name = self._get_collection_name(kb_name)
        
        if not self.collection_exists(collection_name):
            collections_result = self.list_collections()
            available_kbs = [info["kb_name"] for info in collections_result.get("collections", [])]
            return {
                "success": False,
                "message": f"Knowledge base '{kb_name}' does not exist. Please upload documents first.",
                "suggestion": f"Try: upload_document_to_kb(kb_name='{kb_name}', ...) to create KB and add documents, then use chat_with_kb.",
                "available_kbs": available_kbs
            }
        
        try:
            # Get vector store
            vector_store = Qdrant(
                client=self.qdrant_client,
                collection_name=collection_name,
                embeddings=self.embed_model
            )
            
            # Get or create memory
            memory = self._get_or_create_memory(collection_name, session_id)
            
            # Create conversational chain
            qa_chain = ConversationalRetrievalChain.from_llm(
                llm=self.llm,
                retriever=vector_store.as_retriever(search_kwargs={"k": top_k}),
                memory=memory,
                return_source_documents=True,
                verbose=False
            )
            
            # Query
            result = qa_chain({"question": query})
            
            # Format sources
            sources = []
            for doc in result.get("source_documents", []):
                sources.append({
                    "content": doc.page_content[:200] + "...",
                    "metadata": doc.metadata
                })
            
            return {
                "success": True,
                "kb_name": kb_name,
                "session_id": session_id,
                "answer": result["answer"],
                "sources": sources
            }
        except Exception as e:
            logger.error(f"Chat failed: {e}", exc_info=True)
            import traceback
            pass  # traceback handled by logger
            return {
                "success": False,
                "message": str(e)
            }
    
    def route_to_kb(self, query: str) -> Optional[Tuple[str, float]]:
        """
        Route query to the most relevant KB using semantic search on master router index
        
        Args:
            query: User query
            
        Returns:
            Tuple of (kb_name, similarity_score) or None if no good match found
        """
        try:
            # Check if router index exists and has data
            if not self.collection_exists(self.ROUTER_COLLECTION_NAME):
                logger.warning("Router index does not exist")
                return None
            
            collection_info = self.qdrant_client.get_collection(self.ROUTER_COLLECTION_NAME)
            if collection_info.points_count == 0:
                logger.warning("Router index is empty (no KBs registered)")
                return None
            
            # Generate query embedding
            query_vector = self.embed_model.embed_query(query)
            
            # Search router index
            search_result = self.qdrant_client.search(
                collection_name=self.ROUTER_COLLECTION_NAME,
                query_vector=query_vector,
                limit=1
            )
            
            if not search_result:
                logger.warning("No results from router index")
                return None
            
            # Get best match
            best_match = search_result[0]
            kb_name = best_match.payload.get("kb_name")
            score = best_match.score
            description = best_match.payload.get("description", "")
            
            logger.info(f"🎯 Router found: {kb_name} (score: {score:.3f})")
            logger.debug(f"Description: {description}")
            
            # Check if score meets threshold
            if score < self.ROUTER_SIMILARITY_THRESHOLD:
                logger.warning(f"Score {score:.3f} below threshold {self.ROUTER_SIMILARITY_THRESHOLD}")
                return None
            
            return (kb_name, score)
            
        except Exception as e:
            logger.error(f"Routing failed: {e}", exc_info=True)
            import traceback
            pass  # traceback handled by logger
            return None
    
    def chat_auto_route(
        self,
        query: str,
        session_id: str,
        top_k: int = 5
    ) -> Dict[str, Any]:
        """
        Chat with automatic KB routing (Semantic Router)
        The system will automatically find the most relevant KB for your query
        
        Args:
            query: User query
            session_id: Session ID for conversation history
            top_k: Number of documents to retrieve
            
        Returns:
            Dict with answer and routing information
        """
        logger.info(f"🌐 Auto-routing query: '{query[:100]}...'")
        
        # Route to best KB
        routing_result = self.route_to_kb(query)
        
        if routing_result is None:
            # No suitable KB found
            available_kbs = self.list_collections()
            kb_list = [kb['kb_name'] for kb in available_kbs.get('collections', [])]
            
            return {
                "success": False,
                "message": "I don't know which knowledge base to use for this question.",
                "suggestion": "Please specify a knowledge base explicitly using chat_with_kb, or upload relevant documents first.",
                "query": query,
                "available_kbs": kb_list,
                "routing_attempted": True,
                "routing_failed": True
            }
        
        # Unpack routing result
        kb_name, confidence_score = routing_result
        
        logger.info(f"✅ Routed to: {kb_name} (confidence: {confidence_score:.3f})")
        
        # Chat with the selected KB
        result = self.chat_with_collection(
            kb_name=kb_name,
            query=query,
            session_id=session_id,
            top_k=top_k
        )
        
        # Add routing information to response
        if result.get("success"):
            result["routed_to"] = kb_name
            result["routing_confidence"] = confidence_score
            result["routing_method"] = "semantic_similarity"
            logger.info(f"✅ Auto-route successful: {kb_name}")
        
        return result
    
    def clear_chat_history(self, kb_name: str, session_id: str) -> Dict[str, Any]:
        """Clear chat history for a session"""
        collection_name = self._get_collection_name(kb_name)
        
        if collection_name in self.chat_histories:
            if session_id in self.chat_histories[collection_name]:
                del self.chat_histories[collection_name][session_id]
                return {
                    "success": True,
                    "message": f"Chat history cleared for session {session_id}"
                }
        
        return {
            "success": False,
            "message": "Session not found"
        }


# Global instance
_multi_kb_rag = None

def get_multi_kb_rag() -> MultiKnowledgeBaseRAG:
    """Get global Multi-KB RAG instance (singleton)"""
    global _multi_kb_rag
    if _multi_kb_rag is None:
        _multi_kb_rag = MultiKnowledgeBaseRAG()
    return _multi_kb_rag
