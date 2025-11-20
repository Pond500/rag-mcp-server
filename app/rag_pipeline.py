# app/rag_pipeline.py (เวอร์ชัน LangChain)

import app.config as config
from app.ocr_service import get_text_from_pdf
import os
import json
import qdrant_client
from qdrant_client.models import Distance, VectorParams

# --- LangChain Core ---
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import Runnable, RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_core.callbacks import CallbackManagerForRetrieverRun

# --- Models (LLM, Embedding) ---
from langchain_openai import ChatOpenAI
from langchain_huggingface import HuggingFaceEmbeddings

# --- Text Splitter ---
from langchain_text_splitters import RecursiveCharacterTextSplitter

# --- Vector DB (Qdrant) ---
from langchain_qdrant import Qdrant

# --- Reranker ---
from langchain.retrievers.document_compressors import CrossEncoderReranker
from langchain_community.cross_encoders import HuggingFaceCrossEncoder
from langchain.retrievers import ContextualCompressionRetriever

# --- Chains ---
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain

# --- Python ---
from typing import List, Dict, Any, Tuple, Optional
import io

# ---------------------------------------------------------------------
# 1. GLOBAL MODELS & SETTINGS
# (แทนที่ setup_global_settings() เดิม)
# ---------------------------------------------------------------------
print("Setting up LangChain Global Models...")

# 1.1. สร้าง LLM
llm = ChatOpenAI(
    model=config.LLM_MODEL_NAME,
    base_url=config.LLM_API_BASE,
    api_key=config.LLM_API_KEY,
    max_tokens=1500,
    temperature=0.5,
    request_timeout=120
)

# 1.2. สร้าง Embedding Model
embed_model = HuggingFaceEmbeddings(
    model_name=config.EMBED_MODEL_NAME,
    # (bge-m3 แนะนำให้ normalize)
    encode_kwargs={'normalize_embeddings': True}
)

# 1.3. สร้าง Text Splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=config.CHUNK_SIZE,
    chunk_overlap=config.CHUNK_OVERLAP
)

print("Global Models setup complete.")

# ---------------------------------------------------------------------
# 2. VECTOR STORE LOADER
# ---------------------------------------------------------------------

def get_vector_store() -> Qdrant:
    """เชื่อมต่อ Qdrant Client และสร้าง collection ถ้ายังไม่มี"""
    client = qdrant_client.QdrantClient(
        host=config.QDRANT_HOST,
        port=config.QDRANT_PORT
    )
    
    # ตรวจสอบว่า collection มีอยู่หรือไม่ ถ้าไม่มีให้สร้างใหม่
    try:
        client.get_collection(config.QDRANT_COLLECTION_NAME)
        print(f"  Collection '{config.QDRANT_COLLECTION_NAME}' already exists.")
    except Exception:
        print(f"  Creating new collection: '{config.QDRANT_COLLECTION_NAME}'...")
        # สร้าง collection ใหม่ (ต้องระบุ vector size ตาม embedding model)
        # BAAI/bge-m3 มี dimension = 1024
        client.create_collection(
            collection_name=config.QDRANT_COLLECTION_NAME,
            vectors_config=VectorParams(size=1024, distance=Distance.COSINE)
        )
        print(f"  Collection created successfully.")
    
    # LangChain Qdrant object จะใช้สำหรับทั้งการเพิ่ม (add) และค้นหา (retrieve)
    return Qdrant(
        client=client,
        collection_name=config.QDRANT_COLLECTION_NAME,
        embeddings=embed_model # <-- (สำคัญ) ต้องระบุ embedding model
    )

def _extract_metadata_from_text(text: str) -> Dict[str, Any]:
    """
    (อัปเดต) เรียก LLM เพื่อสกัด metadata (doc_type, category, etc.) จากข้อความ
    """
    print("  Extracting metadata using LangChain LLM...")
    
    # จำกัดข้อความที่ส่ง (เช่น 4000 ตัวอักษรแรก) เพื่อประหยัด token
    truncated_text = text[:4000]
    
    try:
        # (ยังคงใช้ Prompt เดิมจาก app.prompts ได้)
        from app.prompts import METADATA_EXTRACTOR_TEMPLATE
        print(f"  [DEBUG] Prompt template loaded")
        
        # สร้าง prompt
        print(f"  [DEBUG] Formatting prompt with text length: {len(truncated_text)}")
        prompt_str = METADATA_EXTRACTOR_TEMPLATE.format(context_str=truncated_text)
        print(f"  [DEBUG] Prompt formatted successfully")
        print(f"  [DEBUG] Calling LLM for metadata extraction...")
        
        # (NEW) เรียก LLM (ใช้ .invoke() แทน .complete())
        response = llm.invoke(prompt_str)
        print(f"  [DEBUG] LLM response received")
        
        # (NEW) ผลลัพธ์จะอยู่ใน .content
        raw_output = response.content 
        
        # พยายาม parse JSON จาก LLM output
        json_str = raw_output.strip().strip("```json").strip("```")
        
        extracted_data = json.loads(json_str)
        
        if isinstance(extracted_data, dict):
            print(f"  Successfully extracted metadata: {extracted_data}")
            return extracted_data
        else:
            raise json.JSONDecodeError("LLM did not return a dictionary", "", 0)

    except Exception as e:
        print(f"  [Error] Failed to extract metadata: {type(e).__name__}: {e}")
        import traceback
        print(traceback.format_exc())
        return {
            "doc_type": "Unknown", "category": "Unknown", "status": "Unknown", "title": "N/A"
        }

# ---------------------------------------------------------------------
# 3. INDEXING PIPELINE (อัปเดต - รองรับหลายประเภทไฟล์)
# ---------------------------------------------------------------------

def _get_text_from_txt(file_bytes: bytes) -> List[Dict[str, Any]]:
    """ดึงข้อความจากไฟล์ TXT"""
    try:
        text = file_bytes.decode('utf-8')
        return [{"page_number": 1, "text": text}]
    except Exception as e:
        print(f"[Error] Failed to read TXT file: {e}")
        return []

def _get_text_from_docx(file_bytes: bytes) -> List[Dict[str, Any]]:
    """ดึงข้อความจากไฟล์ DOCX"""
    try:
        import docx
        print(f"  [DEBUG] docx module path: {docx.__file__}")
        print(f"  [DEBUG] Creating document from bytes (size: {len(file_bytes)})")
        
        doc = docx.Document(io.BytesIO(file_bytes))
        print(f"  [DEBUG] Document created successfully")
        
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        text = '\n'.join(full_text)
        print(f"  [DEBUG] Extracted {len(full_text)} paragraphs, total length: {len(text)}")
        return [{"page_number": 1, "text": text}]
    except Exception as e:
        print(f"[Error] Failed to read DOCX file: {e}")
        import traceback
        print(traceback.format_exc())
        return []

def index_document(file_bytes: bytes, file_name: str, content_type: str) -> Tuple[bool, Optional[Dict[str, Any]]]:
    """
    (อัปเดต) Pipeline สำหรับ Indexing ด้วย LangChain รองรับ PDF, TXT, DOCX
    """
    print(f"Indexing started for: {file_name} (type: {content_type})")
    
    # 1. ดึงข้อความตามประเภทไฟล์
    if content_type == "application/pdf":
        page_data_list = get_text_from_pdf(file_bytes)
    elif content_type == "text/plain":
        page_data_list = _get_text_from_txt(file_bytes)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        page_data_list = _get_text_from_docx(file_bytes)
    else:
        print(f"[Error] Unsupported file type: {content_type}")
        return False, None
    
    if not page_data_list:
        print(f"[Error] Failed to extract text from {file_name}.")
        return False, None
    
    # 2. Extract Metadata (LLM)
    first_page_text = page_data_list[0].get('text', '')
    if not first_page_text:
        print(f"[Warning] First page has no text. Skipping metadata extraction.")
        extracted_metadata = {
            "doc_type": "Unknown", "category": "Unknown", "status": "Unknown", "title": "N/A"
        }
    else:
        extracted_metadata = _extract_metadata_from_text(first_page_text)

    # 3. Create LangChain Documents
    print(f"Creating LangChain Documents from {len(page_data_list)} pages...")
    documents: List[Document] = []
    
    for page_data in page_data_list:
        doc_metadata = extracted_metadata.copy() 
        doc_metadata.update({
            "file_name": file_name,
            "page_number": page_data['page_number']
        })
        
        doc = Document(
            page_content=page_data['text'],
            metadata=doc_metadata
        )
        documents.append(doc)
    
    # 4. Split Documents
    print(f"Splitting {len(documents)} documents into chunks...")
    split_docs = text_splitter.split_documents(documents)
    print(f"Total chunks created: {len(split_docs)}")
    
    # 5. Index to Qdrant
    print("Connecting to Qdrant Vector Store for indexing...")
    try:
        vector_store = get_vector_store()
        vector_store.add_documents(split_docs)
        
        print(f"Successfully indexed: {file_name}")
        return True, extracted_metadata
    except Exception as e:
        print(f"[Error] Indexing failed for {file_name}: {e}")
        return False, None

# เก็บ index_pdf เดิมไว้เพื่อ backward compatibility
def index_pdf(pdf_bytes: bytes, file_name: str) -> Tuple[bool, Optional[Dict[str, Any]]]:
    """(Deprecated) ใช้ index_document แทน"""
    return index_document(pdf_bytes, file_name, "application/pdf")

# ---------------------------------------------------------------------
# 4. QUERYING PIPELINE (อัปเดต)
# ---------------------------------------------------------------------

# (สำคัญ) เราต้องสร้าง Prompt Template ใหม่ที่เข้ากันได้กับ LangChain Chains
# (LCEL chains คาดหวัง input_variables ชื่อ 'context' และ 'input')
LANGCHAIN_QA_PROMPT_TMPL = (
    "เรามีข้อมูลบริบทดังต่อไปนี้: \n"
    "---------------------\n"
    "{context}\n"
    "---------------------\n"
    "จงใช้ข้อมูลบริบทนี้เพื่อตอบคำถามเท่านั้น ห้ามใช้ความรู้เดิมที่มี"
    "จงตอบอย่างสุภาพและเป็นมิตรเสมอ \n"
    "คำถาม: {input}\n"
    "คำตอบ: "
)
THAI_QA_TEMPLATE_LC = PromptTemplate(
    template=LANGCHAIN_QA_PROMPT_TMPL,
    input_variables=["context", "input"]
)

# Custom Reranker Class ที่เก็บ score
class ScoredCrossEncoderReranker(CrossEncoderReranker):
    """CrossEncoderReranker ที่เพิ่ม score ลงใน metadata"""
    
    def compress_documents(
        self,
        documents: List[Document],
        query: str,
        callbacks: Optional[CallbackManagerForRetrieverRun] = None,
    ) -> List[Document]:
        """Rerank documents และเพิ่ม score ลงใน metadata"""
        
        # คำนวณ score ก่อน (ใช้ score method ของ HuggingFaceCrossEncoder)
        pairs = [(query, doc.page_content) for doc in documents]
        scores = list(self.model.score(pairs))
        
        # จับคู่ docs กับ scores
        docs_with_scores = list(zip(documents, scores))
        
        # Sort by score (descending)
        docs_with_scores.sort(key=lambda x: x[1], reverse=True)
        
        # เลือก top_n
        top_docs = docs_with_scores[:self.top_n]
        
        # เพิ่ม score ลงใน metadata
        reranked_docs = []
        for doc, score in top_docs:
            doc.metadata["relevance_score"] = float(score)
            reranked_docs.append(doc)
        
        return reranked_docs

def get_query_engine() -> Runnable:
    """
    (อัปเดต) สร้าง RAG Query Engine ด้วย LangChain (LCEL)
    """
    print("Building LangChain RAG Query Engine...")
    
    # 1. เชื่อมต่อ Vector Store
    vector_store = get_vector_store()
    
    # 2. สร้าง Retriever (Base)
    retriever = vector_store.as_retriever(
        search_type="similarity",
        search_kwargs={"k": 10}
    )
    
    # 3. สร้าง Reranker (ใช้ Custom Class)
    model = HuggingFaceCrossEncoder(model_name="BAAI/bge-reranker-v2-m3")
    compressor = ScoredCrossEncoderReranker(model=model, top_n=3)
    
    # 4. สร้าง Contextual Compression Retriever (Retriever + Reranker)
    # (นี่คือ Retriever ตัวจริงที่เราจะใช้)
    
    compression_retriever = ContextualCompressionRetriever(
        base_compressor=compressor,
        base_retriever=retriever
    )
    
    # 5. สร้าง RAG Chain (Stateless)
    # (ใช้ helper `create_retrieval_chain` เพื่อให้ได้ cả 'answer' และ 'context' (source_nodes) กลับมา)
    
    # 5.1 Chain ย่อย: สำหรับยัด context + input เข้าไปใน Prompt
    qa_document_chain = create_stuff_documents_chain(llm, THAI_QA_TEMPLATE_LC)
    
    # 5.2 Chain หลัก: (Question) -> Retriever -> (Context) -> QA_Document_Chain -> (Answer)
    # (Chain นี้จะคืนค่า Dict: {'answer': ..., 'context': [...]} )
    rag_chain = create_retrieval_chain(
        compression_retriever,  # (Retriever ที่มี Reranker)
        qa_document_chain
    )
    
    print("LangChain RAG Query Engine built successfully.")
    
    # คืนค่า RAG chain ทั้งก้อน
    return rag_chain